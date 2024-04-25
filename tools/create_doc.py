from docx import Document
from docx.shared import Pt

class CreateDocx:
    def __init__(self, contents):
        self.contents = contents

# Create a new Document
doc = Document()

# Add a heading with different levels
doc.add_heading('Heading 1', level=1)
doc.add_heading('Heading 2', level=2)

# Add text with different formatting
p = doc.add_paragraph()
run = p.add_run('Bold and Italic Text')
run.bold = True
run.italic = True

# Add text with specific font size
p = doc.add_paragraph()
run = p.add_run('Text with specific font size Text with specific font size Text with specific font size Text with specific font size Text with specific font size Text with specific font size Text with specific font size Text with specific font size Text with specific font size Text with specific font size Text with specific font size Text with specific font size Text with specific font size Text with specific font size Text with specific font size Text with specific font size Text with specific font size Text with specific font size Text with specific font size Text with specific font size Text with specific font size Text with specific font size Text with specific font size Text with specific font size Text with specific font size Text with specific font size Text with specific font size Text with specific font size ')
run.font.size = Pt(12)  # Font size in points

# Save the document
doc.save('example.docx')

print('DOCX created: example.docx')